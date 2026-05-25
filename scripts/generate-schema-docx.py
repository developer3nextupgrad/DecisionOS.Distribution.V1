"""Generate DecisionOS Distribution database schema reference DOCX."""
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "_client_docs" / "DecisionOS_Database_Schema_Reference.docx"

# Category tags
IMPORT = "Import"
DASH = "Dashboard"
ADMIN = "Admin / configuration"
AUTH = "Authentication"
BOTH = "Import + Dashboard"

TABLES = [
    {
        "name": "Tenants",
        "purpose": "Client (distributor) identity. Links each tenant to a business profile for KPI/driver library resolution.",
        "category": BOTH,
        "used_by": "Import resolves tenant by ClientId; Index/Dashboard/API select tenant; upload batches scoped by TenantId.",
        "fields": [
            ("Id", "uuid", "PK", "Internal tenant key."),
            ("ClientId", "text", "unique", "External client code (e.g. DIST-001). Used in URLs and CSV import."),
            ("Name", "text", "", "Display name on dashboard and tenant picker."),
            ("Archetype", "text?", "", "Optional business archetype label shown in API."),
            ("BusinessProfileId", "uuid?", "FK → BusinessProfiles", "Which KPI/driver/influencer library applies."),
        ],
    },
    {
        "name": "VerticalLibraries",
        "purpose": "Top-level vertical taxonomy (e.g. distribution vertical packs).",
        "category": ADMIN,
        "used_by": "Admin profile setup; referenced by BusinessProfiles.",
        "fields": [
            ("Id", "uuid", "PK", ""),
            ("Code", "text", "unique", "Short code for the vertical library."),
            ("Name", "text", "", "Display name."),
            ("Description", "text?", "", ""),
            ("IsActive", "boolean", "", "Whether library is selectable."),
        ],
    },
    {
        "name": "BusinessProfiles",
        "purpose": "Reusable KPI/driver template for a class of distributors.",
        "category": ADMIN,
        "used_by": "Admin; tenants inherit definitions via BusinessProfileId.",
        "fields": [
            ("Id", "uuid", "PK", ""),
            ("Code", "text", "unique", "Profile code (e.g. default distribution pack)."),
            ("Name", "text", "", ""),
            ("Description", "text?", "", ""),
            ("VerticalLibraryId", "uuid?", "FK", "Parent vertical."),
            ("ChannelStructure", "text?", "", "Metadata for channel modeling."),
            ("LocationStructure", "text?", "", "Metadata for location hierarchy."),
            ("ActiveKpiProfileCode", "text?", "", "Which KPI set variant is active."),
            ("ThresholdProfileCode", "text?", "", "Threshold pack reference."),
            ("IsActive", "boolean", "", ""),
        ],
    },
    {
        "name": "KpiDefinitions",
        "purpose": "Master KPI catalog: codes, thresholds, units, recommended actions.",
        "category": ADMIN,
        "used_by": "Import/console and upload scoring resolve KPI by Code; dashboard cards join via KpiDefinitionId.",
        "fields": [
            ("Id", "integer", "PK", ""),
            ("BusinessProfileId", "uuid?", "FK", "Null = global; else profile-specific definition."),
            ("Code", "text", "unique per profile", "Pillar/KPI code (e.g. GrossMargin%, DOH, CCC)."),
            ("Name", "text", "", "Human-readable KPI name on dashboard."),
            ("Unit", "text", "", "pct | days | currency-style (formatted in UI)."),
            ("Direction", "integer", "", "KpiDirection enum: higher-is-better vs lower-is-better for status."),
            ("Target", "numeric", "", "Green target value."),
            ("AmberThreshold", "numeric", "", "Boundary for YELLOW status."),
            ("RedThreshold", "numeric", "", "Boundary for RED status."),
            ("MinValue", "numeric?", "", "Optional validation floor."),
            ("MaxValue", "numeric?", "", "Optional validation ceiling."),
            ("AlertPriority", "integer", "", "Weight when selecting top alert."),
            ("RecommendedAction", "text", "", "Default action text for weekly focus."),
            ("DiagnosticChecks", "text", "", "Diagnostic hints (often JSON/text)."),
        ],
    },
    {
        "name": "DriverDefinitions",
        "purpose": "Catalog of operational drivers per pillar (KPI code).",
        "category": ADMIN,
        "used_by": "Validates driver CSV import; optional ranking metadata.",
        "fields": [
            ("Id", "integer", "PK", ""),
            ("BusinessProfileId", "uuid?", "FK", ""),
            ("PillarCode", "text", "", "Must match a KpiDefinitions.Code."),
            ("DriverCode", "text", "", "Stable driver identifier."),
            ("DisplayName", "text", "", ""),
            ("Description", "text?", "", ""),
            ("SortOrder", "integer", "", "Display ordering."),
            ("IsActive", "boolean", "", ""),
        ],
    },
    {
        "name": "InfluencerDefinitions",
        "purpose": "Weighted influencers linked to drivers for advanced analytics.",
        "category": ADMIN,
        "used_by": "Profile seeding; not directly shown on main dashboard today.",
        "fields": [
            ("Id", "integer", "PK", ""),
            ("BusinessProfileId", "uuid?", "FK", ""),
            ("PillarCode", "text", "", ""),
            ("DriverCode", "text", "", ""),
            ("InfluencerCode", "text", "", ""),
            ("DisplayName", "text", "", ""),
            ("Description", "text?", "", ""),
            ("Direction", "integer", "", "InfluencerImpactDirection enum."),
            ("Weight", "integer", "", "Relative influence weight."),
            ("IsActive", "boolean", "", ""),
        ],
    },
    {
        "name": "TenantKpiOverrides",
        "purpose": "Per-tenant overrides of KPI thresholds and metadata.",
        "category": ADMIN,
        "used_by": "Admin overrides UI; status computation uses merged definition.",
        "fields": [
            ("Id", "integer", "PK", ""),
            ("TenantId", "uuid", "FK", ""),
            ("KpiCode", "text", "unique per tenant", ""),
            ("Target", "numeric?", "", ""),
            ("AmberThreshold", "numeric?", "", ""),
            ("RedThreshold", "numeric?", "", ""),
            ("MinValue", "numeric?", "", ""),
            ("MaxValue", "numeric?", "", ""),
            ("AlertPriority", "integer?", "", ""),
            ("RecommendedAction", "text?", "", ""),
            ("DiagnosticChecks", "text?", "", ""),
            ("IsActive", "boolean", "", "Disable KPI for tenant when false."),
        ],
    },
    {
        "name": "TenantDriverOverrides",
        "purpose": "Per-tenant overrides of driver display/active flags.",
        "category": ADMIN,
        "used_by": "Admin overrides UI.",
        "fields": [
            ("Id", "integer", "PK", ""),
            ("TenantId", "uuid", "FK", ""),
            ("PillarCode", "text", "", ""),
            ("DriverCode", "text", "", ""),
            ("DisplayName", "text?", "", ""),
            ("Description", "text?", "", ""),
            ("SortOrder", "integer?", "", ""),
            ("IsActive", "boolean?", "", ""),
        ],
    },
    {
        "name": "UploadBatches",
        "purpose": "One weekly upload package per tenant/period (Operations UI + batch import).",
        "category": IMPORT,
        "used_by": "Operations/Uploads pages; UploadBatchImportService; links to ImportRun after commit.",
        "fields": [
            ("Id", "bigint", "PK", ""),
            ("TenantId", "uuid", "FK", ""),
            ("PeriodEnd", "date", "", "Reporting week end date for the package."),
            ("Status", "text", "", "Workflow status (e.g. Draft, Validated, Imported)."),
            ("ReadinessStatus", "text?", "", "ReadyToRun | ReadyWithLimitations | NotReady — drives DataConfidence."),
            ("ValidationSummary", "text?", "", "Human-readable validation rollup."),
            ("CreatedAt", "timestamptz", "", ""),
            ("ImportRunId", "integer?", "FK → ImportRuns", "Set when batch is promoted to a completed import run."),
        ],
    },
    {
        "name": "UploadedFiles",
        "purpose": "Metadata for each CSV/file in an upload batch.",
        "category": IMPORT,
        "used_by": "File storage under App_Data/uploads; mapping and normalization.",
        "fields": [
            ("Id", "bigint", "PK", ""),
            ("UploadBatchId", "bigint", "FK", ""),
            ("ReportType", "integer", "", "ReportType enum: Sales, Inventory, AR, AP, etc."),
            ("OriginalFileName", "text", "", "User-facing filename."),
            ("StoredFileName", "text", "", "Unique stored name."),
            ("StoredRelativePath", "text", "", "Path under App_Data/uploads."),
            ("Sha256Hex", "text", "indexed", "Dedup / integrity fingerprint."),
            ("HeaderRowNumber", "integer", "", "Which row contains CSV headers."),
            ("PeriodStart", "date?", "", "Optional period from file metadata."),
            ("PeriodEnd", "date?", "", "Optional period from file."),
            ("SnapshotDate", "date?", "", "Point-in-time for AR/AP/Inventory snapshots."),
            ("UploadedAt", "timestamptz", "", ""),
        ],
    },
    {
        "name": "UploadedFileColumnMaps",
        "purpose": "Per-file mapping from source CSV column to system field.",
        "category": IMPORT,
        "used_by": "Operations Map page; normalization reads maps.",
        "fields": [
            ("Id", "bigint", "PK", ""),
            ("UploadedFileId", "bigint", "FK", ""),
            ("SourceColumn", "text", "unique per file", "CSV header name."),
            ("SystemField", "text?", "", "SystemFields name (e.g. Net_Sales, Open_Balance)."),
            ("Ignore", "boolean", "", "Skip column when true."),
        ],
    },
    {
        "name": "MappingTemplates",
        "purpose": "Reusable column-mapping templates per tenant and report type.",
        "category": IMPORT,
        "used_by": "Operations mapping UI to pre-fill column maps.",
        "fields": [
            ("Id", "bigint", "PK", ""),
            ("TenantId", "uuid", "FK", ""),
            ("ReportType", "integer", "", "ReportType enum."),
            ("Name", "text", "unique per tenant+type", "Template label."),
            ("CreatedAt", "timestamptz", "", ""),
        ],
    },
    {
        "name": "MappingRules",
        "purpose": "Column rules belonging to a mapping template.",
        "category": IMPORT,
        "used_by": "Applied when user selects a template.",
        "fields": [
            ("Id", "bigint", "PK", ""),
            ("MappingTemplateId", "bigint", "FK", ""),
            ("SourceColumn", "text", "", ""),
            ("SystemField", "text?", "", ""),
            ("Ignore", "boolean", "", ""),
        ],
    },
    {
        "name": "UploadBatchIssues",
        "purpose": "Validation issues for an upload batch (missing mappings, package gaps).",
        "category": IMPORT,
        "used_by": "ValidateAsync in UploadBatchImportService; Operations details UI.",
        "fields": [
            ("Id", "bigint", "PK", ""),
            ("UploadBatchId", "bigint", "FK", ""),
            ("UploadedFileId", "bigint?", "FK", "Null = batch-level issue."),
            ("Severity", "integer", "", "Critical / Warning severity enum."),
            ("Category", "text", "", "File | Mapping | RequiredField | Package, etc."),
            ("Field", "text?", "", "Related system field name."),
            ("Message", "text", "", "User-facing issue text."),
        ],
    },
    {
        "name": "NormalizedSalesRows",
        "purpose": "Staging: normalized sales transaction lines from uploaded Sales CSV.",
        "category": IMPORT,
        "used_by": "UploadBatchImportService computes GrossMargin%, DOH, CCC inputs.",
        "fields": [
            ("Id", "bigint", "PK", ""),
            ("TenantId", "uuid", "", ""),
            ("PeriodEnd", "date", "", "Batch reporting period."),
            ("UploadBatchId", "bigint", "FK", ""),
            ("UploadedFileId", "bigint", "FK", ""),
            ("SourceRowNumber", "integer", "", "Original CSV row for traceability."),
            ("Status", "integer", "", "RowStatus: Valid, Warning, Rejected, etc."),
            ("IssueSummary", "text?", "", "Row-level validation message."),
            ("RawJson", "text", "", "Full mapped row as JSON backup."),
            ("TransactionDate", "date?", "", "From Transaction_Date mapping."),
            ("TransactionId", "text?", "", ""),
            ("CustomerId", "text?", "", ""),
            ("CustomerName", "text?", "", ""),
            ("SkuId", "text?", "", ""),
            ("ProductDescription", "text?", "", ""),
            ("LocationId", "text?", "", ""),
            ("QuantitySold", "numeric?", "", ""),
            ("NetSales", "numeric?", "", "Primary revenue input for KPIs."),
            ("Cogs", "numeric?", "", "Cost of goods for margin and CCC."),
            ("GrossProfit", "numeric?", "", "Optional; may be derived."),
        ],
    },
    {
        "name": "NormalizedInventoryRows",
        "purpose": "Staging: inventory snapshot lines for DOH and inventory value KPIs.",
        "category": IMPORT,
        "used_by": "UploadBatchImportService DoH calculation.",
        "fields": [
            ("Id", "bigint", "PK", ""),
            ("TenantId", "uuid", "", ""),
            ("PeriodEnd", "date", "", ""),
            ("UploadBatchId", "bigint", "FK", ""),
            ("UploadedFileId", "bigint", "FK", ""),
            ("SourceRowNumber", "integer", "", ""),
            ("Status", "integer", "", "RowStatus enum."),
            ("IssueSummary", "text?", "", ""),
            ("RawJson", "text", "", ""),
            ("SnapshotDate", "date?", "", ""),
            ("SkuId", "text?", "", ""),
            ("LocationId", "text?", "", ""),
            ("QuantityOnHand", "numeric?", "", ""),
            ("InventoryValue", "numeric?", "", "Summed for days-on-hand."),
            ("AverageCost", "numeric?", "", ""),
            ("LastSaleDate", "date?", "", ""),
        ],
    },
    {
        "name": "NormalizedArRows",
        "purpose": "Staging: accounts receivable open items.",
        "category": IMPORT,
        "used_by": "AR_PastDue31p% and CCC (DSO) KPI computation.",
        "fields": [
            ("Id", "bigint", "PK", ""),
            ("TenantId", "uuid", "", ""),
            ("PeriodEnd", "date", "", ""),
            ("UploadBatchId", "bigint", "FK", ""),
            ("UploadedFileId", "bigint", "FK", ""),
            ("SourceRowNumber", "integer", "", ""),
            ("Status", "integer", "", ""),
            ("IssueSummary", "text?", "", ""),
            ("RawJson", "text", "", ""),
            ("SnapshotDate", "date?", "", ""),
            ("CustomerId", "text?", "", ""),
            ("CustomerName", "text?", "", ""),
            ("InvoiceId", "text?", "", ""),
            ("DueDate", "date?", "", ""),
            ("OpenBalance", "numeric?", "", ""),
            ("AgingBucket", "text?", "", "Used with DaysPastDue for past-due logic."),
            ("DaysPastDue", "integer?", "", ""),
        ],
    },
    {
        "name": "NormalizedApRows",
        "purpose": "Staging: accounts payable open bills.",
        "category": IMPORT,
        "used_by": "AP_PastDue31p% and CCC (DPO) KPI computation.",
        "fields": [
            ("Id", "bigint", "PK", ""),
            ("TenantId", "uuid", "", ""),
            ("PeriodEnd", "date", "", ""),
            ("UploadBatchId", "bigint", "FK", ""),
            ("UploadedFileId", "bigint", "FK", ""),
            ("SourceRowNumber", "integer", "", ""),
            ("Status", "integer", "", ""),
            ("IssueSummary", "text?", "", ""),
            ("RawJson", "text", "", ""),
            ("SnapshotDate", "date?", "", ""),
            ("VendorId", "text?", "", ""),
            ("VendorName", "text?", "", ""),
            ("BillId", "text?", "", ""),
            ("DueDate", "date?", "", ""),
            ("OpenBalance", "numeric?", "", ""),
            ("AgingBucket", "text?", "", ""),
            ("DaysPastDue", "integer?", "", ""),
        ],
    },
    {
        "name": "ImportRuns",
        "purpose": "Audit record for each KPI/driver import execution (console CSV or batch promotion).",
        "category": IMPORT,
        "used_by": "Console Import app; API returns latest run metadata on week payload.",
        "fields": [
            ("Id", "integer", "PK", ""),
            ("TenantId", "uuid", "FK", ""),
            ("PeriodEnd", "date", "", ""),
            ("StartedAt", "timestamptz", "", ""),
            ("CompletedAt", "timestamptz?", "", ""),
            ("Status", "text", "", "Running | Completed | Failed."),
            ("KpiRowsProcessed", "integer", "", "Rows written from KPI CSV or computed KPIs."),
            ("DriverRowsProcessed", "integer", "", "Driver CSV row count."),
            ("SourceFingerprint", "text?", "", "Hash of source files for idempotent skip."),
            ("ReadinessStatus", "text?", "", "Data quality gate from validation."),
            ("ValidationSummary", "text?", "", ""),
            ("ErrorMessage", "text?", "", "Failure detail when Status=Failed."),
        ],
    },
    {
        "name": "ImportRunIssues",
        "purpose": "Row/field-level issues recorded during console CSV import.",
        "category": IMPORT,
        "used_by": "DecisionOS.Distribution.Import Program.cs validation.",
        "fields": [
            ("Id", "bigint", "PK", ""),
            ("ImportRunId", "integer", "FK", ""),
            ("Severity", "integer", "", "ImportValidationSeverity."),
            ("Category", "text", "", "KPI | Driver | etc."),
            ("RowNumber", "integer?", "", "CSV row number."),
            ("Field", "text?", "", "Column/field name."),
            ("Message", "text", "", ""),
        ],
    },
    {
        "name": "KpiSnapshots",
        "purpose": "Per-tenant, per-week KPI measured values and RAG status — core dashboard fact table.",
        "category": BOTH,
        "used_by": "Written by console KPI CSV import OR UploadBatchImportService; read by Dashboard, Index week list, API.",
        "fields": [
            ("Id", "integer", "PK", ""),
            ("TenantId", "uuid", "FK", ""),
            ("PeriodEnd", "date", "", "Week ending date (unique with KpiDefinitionId)."),
            ("KpiDefinitionId", "integer", "FK", ""),
            ("Value", "numeric", "", "Measured KPI value."),
            ("Status", "text", "", "GREEN | YELLOW | RED | GRAY."),
            ("WeekOverWeekDelta", "numeric?", "", "Change vs prior week (set post-import)."),
            ("DataConfidence", "text?", "", "High | Medium | Low from import readiness."),
            ("CardDetailLine1", "text?", "", "Extra card text (CSV import or insufficient-data message)."),
            ("CardDetailLine2", "text?", "", "Second detail line on KPI card."),
        ],
    },
    {
        "name": "DriverValues",
        "purpose": "Ranked operational drivers under each KPI pillar for a tenant/week.",
        "category": BOTH,
        "used_by": "Console driver CSV import; Dashboard and API driver lists (filtered by top alert pillar).",
        "fields": [
            ("Id", "integer", "PK", ""),
            ("TenantId", "uuid", "FK", ""),
            ("PeriodEnd", "date", "", ""),
            ("PillarCode", "text", "", "Matches KpiDefinitions.Code."),
            ("DriverCode", "text?", "", "Optional stable code from definitions."),
            ("DriverName", "text", "", "Display name."),
            ("Dimension1", "text?", "", "Slice dimension (e.g. category)."),
            ("Dimension2", "text?", "", "Second slice."),
            ("Current", "numeric", "", "Primary metric value."),
            ("WeekOverWeekDelta", "numeric?", "", ""),
            ("Context", "text?", "", "Narrative context."),
            ("Rank", "integer", "", "Sort within pillar (1 = top)."),
            ("Status", "text", "", "GREEN | YELLOW | RED."),
            ("WhyItMatters", "text", "", "Explanation for operators."),
            ("Owner", "text?", "", "Accountable person."),
            ("AssignedSummary", "text?", "", "Holdover view: assigned target narrative."),
            ("TargetSummary", "text?", "", "Holdover view: target narrative."),
            ("CurrentSummary", "text?", "", "Holdover view: current state narrative."),
            ("FixProgressPercent", "integer?", "", "0–100 fix progress; UI falls back from Status."),
        ],
    },
    {
        "name": "Alerts",
        "purpose": "Single top-priority alert for the tenant/week (worst KPI by rules).",
        "category": DASH,
        "used_by": "Generated after import by AlertService; Dashboard hero + API TopAlert; filters drivers.",
        "fields": [
            ("Id", "integer", "PK", ""),
            ("TenantId", "uuid", "FK", ""),
            ("PeriodEnd", "date", "unique per tenant+week", ""),
            ("KpiDefinitionId", "integer", "FK", "Which pillar triggered the alert."),
            ("Severity", "text", "", "Alert severity label."),
            ("ReasonSummary", "text", "", "Why this KPI is the focus."),
        ],
    },
    {
        "name": "WeeklyFocuses",
        "purpose": "One recommended weekly decision focus tied to the top alert.",
        "category": DASH,
        "used_by": "WeeklyFocusService after import; Dashboard and API WeeklyFocus section.",
        "fields": [
            ("Id", "integer", "PK", ""),
            ("TenantId", "uuid", "FK", ""),
            ("PeriodEnd", "date", "unique per tenant+week", ""),
            ("KpiDefinitionId", "integer", "FK", ""),
            ("DecisionQuestion", "text", "", "Question leadership should answer."),
            ("RecommendedAction", "text", "", "Suggested action."),
            ("WhyNow", "text", "", "Urgency narrative."),
            ("Owner", "text", "", "Suggested owner."),
            ("Cadence", "text", "", "Review cadence (e.g. weekly)."),
        ],
    },
    {
        "name": "ActionItems",
        "purpose": "Execution tracking: accountable follow-ups linked to KPI pillars.",
        "category": DASH,
        "used_by": "API /api/tenants/{clientId}/weeks/{periodEnd} Actions; status updates via POST.",
        "fields": [
            ("Id", "bigint", "PK", ""),
            ("TenantId", "uuid", "FK", ""),
            ("PeriodEnd", "date", "", "Week the action belongs to."),
            ("KpiDefinitionId", "integer?", "FK", "Optional link to pillar."),
            ("Title", "text", "", ""),
            ("Description", "text?", "", ""),
            ("Owner", "text", "", ""),
            ("DueDate", "date?", "", ""),
            ("Status", "text", "", "NotStarted | InProgress | AtRisk | Completed | Deferred | Blocked."),
            ("Priority", "integer", "", "Lower = higher priority in lists."),
            ("Notes", "text?", "", ""),
            ("CreatedAt", "timestamptz", "", ""),
            ("UpdatedAt", "timestamptz", "", ""),
            ("CompletedAt", "timestamptz?", "", "Set when status → Completed."),
        ],
    },
]

IDENTITY_TABLES = [
    ("AspNetUsers", AUTH, "ApplicationUser: login identity with DisplayName extension."),
    ("AspNetRoles", AUTH, "Roles: Admin, Operator, Viewer, Developer."),
    ("AspNetUserRoles", AUTH, "User-to-role assignments."),
    ("AspNetUserClaims", AUTH, "Extra claims per user."),
    ("AspNetRoleClaims", AUTH, "Extra claims per role."),
    ("AspNetUserLogins", AUTH, "External login providers."),
    ("AspNetUserTokens", AUTH, "Password reset / 2FA tokens."),
]


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True


def add_field_table(doc, fields):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Column", "Type", "Constraints", "Purpose"]):
        hdr[i].text = h
    for col, typ, constraints, purpose in fields:
        row = table.add_row().cells
        row[0].text = col
        row[1].text = typ
        row[2].text = constraints
        row[3].text = purpose
    doc.add_paragraph()


def build():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_heading(doc, "DecisionOS Distribution — Database Schema Reference", 0)
    add_para(
        doc,
        f"Generated from EF Core model snapshot. PostgreSQL database. Date: {date.today().isoformat()}.",
    )
    doc.add_paragraph()

    add_heading(doc, "1. Overview", 1)
    add_para(
        doc,
        "DecisionOS.Distribution persists data in PostgreSQL via Entity Framework Core 8. "
        "Two import paths feed dashboard metrics: (A) Operations upload batches — CSV files mapped to system fields, "
        "normalized into staging tables, then KPIs computed; (B) Console import — direct KPI and driver CSV files. "
        "Both paths write KpiSnapshots, DriverValues, Alerts, and WeeklyFocuses consumed by the Razor Dashboard and /api endpoints.",
    )

    add_heading(doc, "2. Category legend", 1)
    legend = doc.add_table(rows=1, cols=2)
    legend.style = "Table Grid"
    legend.rows[0].cells[0].text = "Category"
    legend.rows[0].cells[1].text = "Meaning"
    for cat, meaning in [
        (IMPORT, "Upload pipeline, staging, mapping, import audit — not shown on main dashboard UI."),
        (DASH, "Read by Dashboard.cshtml and/or /api week payload for end-user decision views."),
        (BOTH, "Written during import and read by dashboard/API."),
        (ADMIN, "Admin UI and definition resolution; configures how import scores KPIs."),
        (AUTH, "ASP.NET Core Identity only."),
    ]:
        r = legend.add_row().cells
        r[0].text = cat
        r[1].text = meaning
    doc.add_paragraph()

    add_heading(doc, "3. Import data flow", 1)
    add_para(doc, "Upload path: UploadBatches → UploadedFiles → UploadedFileColumnMaps → Normalized*Rows → KpiSnapshots (+ Alerts, WeeklyFocuses).")
    add_para(doc, "Console path: ImportRuns ← KPI CSV → KpiSnapshots; optional Driver CSV → DriverValues; then WoW delta, Alerts, WeeklyFocuses.")
    add_para(doc, "UploadBatches.ImportRunId links a committed batch to its ImportRuns audit row.")

    add_heading(doc, "4. Dashboard data flow", 1)
    add_para(
        doc,
        "Index page lists tenants and distinct KpiSnapshots.PeriodEnd weeks. Dashboard loads: Tenants, KpiSnapshots (+ KpiDefinitions), "
        "Alerts (one), DriverValues (top 50), WeeklyFocuses (one). API adds ActionItems and ImportRuns metadata; holdover view widens drivers/actions.",
    )

    add_heading(doc, "5. Application tables", 1)
    for t in TABLES:
        add_heading(doc, t["name"], 2)
        p = doc.add_paragraph()
        p.add_run("Category: ").bold = True
        p.add_run(t["category"])
        doc.add_paragraph(f"Purpose: {t['purpose']}")
        doc.add_paragraph(f"Used by: {t['used_by']}")
        add_field_table(doc, t["fields"])

    add_heading(doc, "6. Identity tables (ASP.NET Core)", 1)
    add_para(doc, "Standard Identity schema with GUID keys. Not part of import or dashboard analytics.")
    it = doc.add_table(rows=1, cols=3)
    it.style = "Table Grid"
    it.rows[0].cells[0].text = "Table"
    it.rows[0].cells[1].text = "Category"
    it.rows[0].cells[2].text = "Notes"
    for name, cat, notes in IDENTITY_TABLES:
        r = it.add_row().cells
        r[0].text = name
        r[1].text = cat
        r[2].text = notes
    doc.add_paragraph()
    add_para(doc, "AspNetUsers extra column: DisplayName (text, optional friendly name).", bold=False)

    add_heading(doc, "7. Quick reference — table by category", 1)
    groups = {IMPORT: [], DASH: [], BOTH: [], ADMIN: [], AUTH: []}
    for t in TABLES:
        groups[t["category"]].append(t["name"])
    for cat in [IMPORT, BOTH, DASH, ADMIN]:
        add_heading(doc, cat, 3)
        doc.add_paragraph(", ".join(sorted(groups[cat])) if groups[cat] else "(none)")
    add_heading(doc, AUTH, 3)
    doc.add_paragraph(", ".join(x[0] for x in IDENTITY_TABLES))

    add_heading(doc, "8. KPI codes computed from uploads (V1)", 1)
    add_para(
        doc,
        "UploadBatchImportService aggregates Normalized* tables into KpiSnapshots for: "
        "GrossMargin%, AR_PastDue31p%, AP_PastDue31p%, DOH, CCC. "
        "NetProfit% and PerfectOrderRate are placeholders (null → GRAY snapshot). "
        "Console CSV import can populate any KPI defined in KpiDefinitions.",
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
